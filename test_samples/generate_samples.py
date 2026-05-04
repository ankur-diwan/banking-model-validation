#!/usr/bin/env python3
"""
Generate sample test files for Model Validation System
"""

import csv
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_model_documentation():
    """Create a sample model documentation DOCX file"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Credit Risk Model Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Executive Summary
    doc.add_heading('1. Executive Summary', 1)
    doc.add_paragraph(
        'This document provides comprehensive documentation for the US Unsecured Personal Loan '
        'Application Scorecard model. The model is designed to predict the probability of default '
        'for new credit applicants in the unsecured personal loan portfolio.'
    )
    
    # Model Overview
    doc.add_heading('2. Model Overview', 1)
    doc.add_paragraph('Model Name: US_Unsecured_Application_Scorecard_v2.1')
    doc.add_paragraph('Model Type: Logistic Regression (GLM)')
    doc.add_paragraph('Product Type: Unsecured Personal Loans')
    doc.add_paragraph('Development Date: January 2024')
    doc.add_paragraph('Model Owner: Credit Risk Analytics Team')
    
    # SR 11-7 Sections
    doc.add_heading('3. SR 11-7 Compliance Framework', 1)
    
    doc.add_heading('3.1 Conceptual Soundness', 2)
    doc.add_paragraph(
        'The model is based on sound economic and statistical principles. '
        'Logistic regression was selected due to its interpretability and proven performance '
        'in credit risk modeling. The model uses 15 predictive variables including:'
    )
    doc.add_paragraph('• Credit Bureau Score (FICO)', style='List Bullet')
    doc.add_paragraph('• Debt-to-Income Ratio', style='List Bullet')
    doc.add_paragraph('• Number of Delinquencies (12 months)', style='List Bullet')
    doc.add_paragraph('• Employment Length', style='List Bullet')
    doc.add_paragraph('• Annual Income', style='List Bullet')
    
    doc.add_heading('3.2 Data Quality and Integrity', 2)
    doc.add_paragraph(
        'Development data consists of 50,000 accounts originated between Jan 2021 and Dec 2022. '
        'Performance window: 12 months. Bad definition: 90+ days past due or charge-off.'
    )
    
    doc.add_heading('3.3 Model Performance', 2)
    doc.add_paragraph('Development Sample Performance:')
    doc.add_paragraph('• KS Statistic: 0.42', style='List Bullet')
    doc.add_paragraph('• Gini Coefficient: 0.68', style='List Bullet')
    doc.add_paragraph('• AUC-ROC: 0.84', style='List Bullet')
    doc.add_paragraph('• Accuracy: 87.5%', style='List Bullet')
    
    doc.add_heading('3.4 Model Assumptions', 2)
    doc.add_paragraph(
        'Key assumptions include: (1) Historical relationships remain stable, '
        '(2) Credit bureau data quality is maintained, (3) Economic conditions remain within '
        'historical ranges, (4) Population characteristics remain consistent.'
    )
    
    doc.add_heading('3.5 Ongoing Monitoring', 2)
    doc.add_paragraph(
        'The model is monitored monthly for: Population Stability (PSI), '
        'Characteristic Stability (CSI), Performance metrics (KS, Gini), '
        'and Discrimination power. Triggers are set at PSI > 0.25 for investigation.'
    )
    
    # Model Variables
    doc.add_heading('4. Model Variables', 1)
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Variable Name'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Coefficient'
    
    # Data rows
    variables = [
        ('Credit Score', 'Continuous', '-0.0045'),
        ('DTI Ratio', 'Continuous', '0.0234'),
        ('Delinquencies_12m', 'Count', '0.3421'),
        ('Employment_Length', 'Categorical', '-0.0876'),
        ('Annual_Income', 'Continuous', '-0.0002'),
    ]
    
    for i, (var, vtype, coef) in enumerate(variables, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = var
        row_cells[1].text = vtype
        row_cells[2].text = coef
    
    # Save
    doc.save('test_samples/sample_model_documentation.docx')
    print("✓ Created: sample_model_documentation.docx")


def create_data_dictionary():
    """Create a sample data dictionary CSV file"""
    with open('test_samples/sample_data_dictionary.csv', 'w', newline='') as f:
        writer = csv.writer(f)
        
        # Header
        writer.writerow([
            'Variable_Name', 'Description', 'Type', 'Valid_Range', 
            'Missing_Treatment', 'Business_Logic'
        ])
        
        # Data rows
        writer.writerow([
            'credit_score', 
            'FICO credit bureau score', 
            'Continuous', 
            '300-850', 
            'Reject application', 
            'Primary risk indicator'
        ])
        writer.writerow([
            'dti_ratio', 
            'Debt-to-Income ratio (%)', 
            'Continuous', 
            '0-100', 
            'Set to median', 
            'Affordability measure'
        ])
        writer.writerow([
            'delinquencies_12m', 
            'Number of 30+ DPD in last 12 months', 
            'Count', 
            '0-20', 
            'Set to 0', 
            'Payment behavior indicator'
        ])
        writer.writerow([
            'employment_length', 
            'Years at current employer', 
            'Continuous', 
            '0-50', 
            'Set to 0', 
            'Stability indicator'
        ])
        writer.writerow([
            'annual_income', 
            'Gross annual income ($)', 
            'Continuous', 
            '10000-500000', 
            'Reject application', 
            'Repayment capacity'
        ])
        writer.writerow([
            'loan_amount', 
            'Requested loan amount ($)', 
            'Continuous', 
            '1000-50000', 
            'N/A', 
            'Exposure amount'
        ])
        writer.writerow([
            'home_ownership', 
            'Housing status', 
            'Categorical', 
            'RENT/OWN/MORTGAGE', 
            'Set to RENT', 
            'Stability indicator'
        ])
        writer.writerow([
            'purpose', 
            'Loan purpose', 
            'Categorical', 
            'DEBT_CONSOLIDATION/HOME/AUTO/OTHER', 
            'Set to OTHER', 
            'Risk segmentation'
        ])
    
    print("✓ Created: sample_data_dictionary.csv")


def create_validation_report():
    """Create a sample validation report DOCX file"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Independent Model Validation Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Model: US Unsecured Application Scorecard v2.1')
    doc.add_paragraph('Validation Date: March 2024')
    doc.add_paragraph('Validator: Model Risk Management Team')
    
    # Executive Summary
    doc.add_heading('Executive Summary', 1)
    doc.add_paragraph(
        'This report presents the findings of the independent validation of the '
        'US Unsecured Application Scorecard model. The validation was conducted in '
        'accordance with SR 11-7 guidelines and internal model risk management policies.'
    )
    
    # Validation Findings
    doc.add_heading('Validation Findings', 1)
    
    doc.add_heading('1. Conceptual Soundness: SATISFACTORY', 2)
    doc.add_paragraph(
        'The model design is appropriate for the intended use. The choice of logistic '
        'regression is well-justified and the variable selection process was rigorous.'
    )
    
    doc.add_heading('2. Data Quality: SATISFACTORY', 2)
    doc.add_paragraph(
        'Development data is representative and of high quality. Missing data treatment '
        'is appropriate. No significant data quality issues identified.'
    )
    
    doc.add_heading('3. Performance: SATISFACTORY', 2)
    doc.add_paragraph('Out-of-time validation results:')
    doc.add_paragraph('• KS Statistic: 0.40 (vs 0.42 development)', style='List Bullet')
    doc.add_paragraph('• Gini: 0.65 (vs 0.68 development)', style='List Bullet')
    doc.add_paragraph('• PSI: 0.08 (Stable)', style='List Bullet')
    
    doc.add_heading('4. Stability: SATISFACTORY', 2)
    doc.add_paragraph(
        'Model shows good stability across time periods and segments. '
        'No significant population shifts detected.'
    )
    
    # Recommendations
    doc.add_heading('Recommendations', 1)
    doc.add_paragraph('1. Implement monthly PSI monitoring', style='List Number')
    doc.add_paragraph('2. Review model annually or when PSI > 0.25', style='List Number')
    doc.add_paragraph('3. Enhance documentation for variable transformations', style='List Number')
    
    # Conclusion
    doc.add_heading('Conclusion', 1)
    doc.add_paragraph(
        'The model is APPROVED for production use with the recommendations noted above. '
        'The model meets all SR 11-7 requirements and is fit for purpose.'
    )
    
    doc.save('test_samples/sample_validation_report.docx')
    print("✓ Created: sample_validation_report.docx")


def create_readme():
    """Create README for test samples"""
    readme_content = """# Test Sample Files for Model Validation System

This directory contains sample files you can use to test the document upload feature.

## Files Included

### 1. sample_model_documentation.docx
**Purpose**: Model development documentation
**Use Case**: Upload this to test document analysis and SR 11-7 section detection
**Contains**:
- Model overview and metadata
- SR 11-7 compliance sections
- Model variables and coefficients
- Performance metrics

### 2. sample_data_dictionary.csv
**Purpose**: Data dictionary for model variables
**Use Case**: Upload to test CSV parsing and variable analysis
**Contains**:
- Variable names and descriptions
- Data types and valid ranges
- Missing value treatment
- Business logic

### 3. sample_validation_report.docx
**Purpose**: Independent validation report
**Use Case**: Upload to test validation report parsing
**Contains**:
- Validation findings
- Performance assessment
- Recommendations
- Approval status

## How to Use

1. **Access the UI**: http://localhost:3002/
2. **Navigate to Step 1**: "Upload Documents (Optional)"
3. **Drag and drop** any of these files or click to browse
4. **Supported formats**: PDF, DOCX, CSV
5. **Continue** to model configuration

## Testing Scenarios

### Scenario 1: Complete Documentation
Upload all three files to test comprehensive document analysis

### Scenario 2: Minimal Upload
Skip document upload and proceed directly to model configuration

### Scenario 3: Single Document
Upload only the model documentation to test partial analysis

## Expected Behavior

- Files are validated for type and size
- Document content is extracted and analyzed
- SR 11-7 sections are automatically detected
- Model information is extracted when available
- Analysis results enhance the validation process

## Notes

- Document upload is **optional** - you can skip this step
- The system generates synthetic data for validation even without uploads
- Uploaded documents provide additional context for validation
- Maximum file size: 10MB per file
- Supported formats: PDF, DOCX, CSV

## Generated by

Banking Model Validation System - Test Sample Generator
"""
    
    with open('test_samples/README.md', 'w') as f:
        f.write(readme_content)
    
    print("✓ Created: README.md")


if __name__ == '__main__':
    print("\n🔧 Generating test sample files...\n")
    
    try:
        create_model_documentation()
        create_data_dictionary()
        create_validation_report()
        create_readme()
        
        print("\n✅ All sample files created successfully!")
        print("\n📁 Files location: test_samples/")
        print("\n📋 Files created:")
        print("   1. sample_model_documentation.docx")
        print("   2. sample_data_dictionary.csv")
        print("   3. sample_validation_report.docx")
        print("   4. README.md")
        print("\n💡 You can now upload these files in the UI at http://localhost:3002/")
        
    except Exception as e:
        print(f"\n❌ Error: {e}")
        print("\nMake sure python-docx is installed:")
        print("   pip install python-docx")

# Made with Bob
