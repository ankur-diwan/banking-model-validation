#!/usr/bin/env python3
"""
Generate Sample Data for Successful Model Validation
Creates synthetic data with good performance metrics that will pass validation
"""

import sys
import os
sys.path.append(os.path.join(os.path.dirname(__file__), '..', 'backend'))

import pandas as pd
import numpy as np
from sklearn.linear_model import LogisticRegression
from sklearn.preprocessing import StandardScaler
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import warnings
warnings.filterwarnings('ignore')

def generate_features(n_samples, seed=42):
    """Generate realistic credit features"""
    np.random.seed(seed)
    
    # Generate correlated features that predict default
    age = np.random.normal(45, 15, n_samples).clip(18, 80)
    income = np.random.lognormal(10.8, 0.6, n_samples).clip(25000, 300000)
    credit_score = np.random.normal(700, 80, n_samples).clip(300, 850)
    
    # DTI ratio - higher means more risk
    dti_ratio = np.random.beta(3, 7, n_samples).clip(0.05, 0.65)
    
    # Delinquencies - count of past delinquencies
    delinquencies = np.random.poisson(0.5, n_samples).clip(0, 10)
    
    # Employment length in months
    employment_months = np.random.exponential(60, n_samples).clip(0, 360)
    
    # Credit utilization
    utilization = np.random.beta(2, 5, n_samples).clip(0, 1)
    
    # Number of credit lines
    num_credit_lines = np.random.poisson(5, n_samples).clip(1, 25)
    
    # Inquiries in last 6 months
    inquiries_6m = np.random.poisson(1, n_samples).clip(0, 10)
    
    df = pd.DataFrame({
        'age': age.astype(int),
        'income': income.round(2),
        'credit_score': credit_score.astype(int),
        'dti_ratio': dti_ratio.round(4),
        'delinquencies': delinquencies.astype(int),
        'employment_months': employment_months.astype(int),
        'utilization': utilization.round(4),
        'num_credit_lines': num_credit_lines.astype(int),
        'inquiries_6m': inquiries_6m.astype(int)
    })
    
    return df

def generate_target_and_score(df, default_rate=0.08, seed=42):
    """
    Generate target variable and scores with good discrimination
    Creates a realistic relationship between features and default
    """
    np.random.seed(seed)
    
    # Create risk score based on features (logit scale)
    risk_score = (
        -0.02 * (df['credit_score'] - 700) +  # Lower credit score = higher risk
        2.5 * df['dti_ratio'] +                # Higher DTI = higher risk
        0.3 * df['delinquencies'] +            # More delinquencies = higher risk
        -0.005 * (df['income'] / 1000) +       # Lower income = higher risk
        0.15 * df['inquiries_6m'] +            # More inquiries = higher risk
        0.5 * df['utilization'] +              # Higher utilization = higher risk
        -0.01 * (df['employment_months'] / 12) # Less employment = higher risk
    )
    
    # Add some noise
    risk_score += np.random.normal(0, 0.5, len(df))
    
    # Convert to probability using logistic function
    probability = 1 / (1 + np.exp(-risk_score))
    
    # Adjust to target default rate
    threshold = np.percentile(probability, (1 - default_rate) * 100)
    adjusted_prob = probability / threshold * default_rate * 2
    adjusted_prob = np.clip(adjusted_prob, 0.001, 0.999)
    
    # Generate target based on probability
    df['target'] = (np.random.random(len(df)) < adjusted_prob).astype(int)
    
    # Convert probability to scorecard scale (300-850, higher = lower risk)
    df['score'] = (850 - (adjusted_prob * 550)).round(0).astype(int)
    
    # Ensure score is inversely related to target
    # Good customers (target=0) should have higher scores
    df.loc[df['target'] == 0, 'score'] = df.loc[df['target'] == 0, 'score'] + np.random.randint(-20, 50, sum(df['target'] == 0))
    df.loc[df['target'] == 1, 'score'] = df.loc[df['target'] == 1, 'score'] + np.random.randint(-50, 20, sum(df['target'] == 1))
    
    # Clip to valid range
    df['score'] = df['score'].clip(300, 850)
    
    return df

def create_model_documentation(filename='successful_model_documentation.docx'):
    """Create comprehensive model documentation with SR 11-7 sections"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Credit Risk Model Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph('US Unsecured Personal Loan Application Scorecard')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    
    doc.add_paragraph()  # Spacing
    
    # Executive Summary
    doc.add_heading('1. Executive Summary', 1)
    doc.add_paragraph(
        'This document provides comprehensive documentation for the US Unsecured Personal Loan '
        'Application Scorecard model (Version 2.1). The model is designed to predict the probability '
        'of default for new credit applicants in the unsecured personal loan portfolio. The model '
        'demonstrates excellent discrimination power with a Gini coefficient of 0.65 and KS statistic '
        'of 0.45 on out-of-sample data.'
    )
    
    # Model Overview
    doc.add_heading('2. Model Overview', 1)
    doc.add_paragraph('Model Name: US_Unsecured_Application_Scorecard_v2.1')
    doc.add_paragraph('Model Type: Logistic Regression (GLM)')
    doc.add_paragraph('Scorecard Type: Application Scorecard')
    doc.add_paragraph('Product Type: Unsecured Personal Loans')
    doc.add_paragraph('Development Date: January 2024')
    doc.add_paragraph('Model Owner: Credit Risk Analytics Team')
    doc.add_paragraph('Model Version: 2.1')
    doc.add_paragraph('Last Validation Date: May 2026')
    
    # SR 11-7 Compliance Framework
    doc.add_heading('3. SR 11-7 Compliance Framework', 1)
    
    doc.add_heading('3.1 Conceptual Soundness', 2)
    doc.add_paragraph(
        'The model is based on sound economic and statistical principles. Logistic regression '
        'was selected due to its interpretability, regulatory acceptance, and proven performance '
        'in credit risk modeling. The model uses 9 predictive variables that have strong theoretical '
        'and empirical relationships with credit default:'
    )
    doc.add_paragraph('• Credit Bureau Score (FICO): Primary indicator of creditworthiness', style='List Bullet')
    doc.add_paragraph('• Debt-to-Income Ratio: Measures repayment capacity', style='List Bullet')
    doc.add_paragraph('• Number of Delinquencies (24 months): Historical payment behavior', style='List Bullet')
    doc.add_paragraph('• Employment Length: Income stability indicator', style='List Bullet')
    doc.add_paragraph('• Annual Income: Repayment capacity', style='List Bullet')
    doc.add_paragraph('• Credit Utilization: Credit management behavior', style='List Bullet')
    doc.add_paragraph('• Number of Credit Lines: Credit experience', style='List Bullet')
    doc.add_paragraph('• Recent Inquiries (6 months): Credit seeking behavior', style='List Bullet')
    doc.add_paragraph('• Age: Proxy for financial maturity', style='List Bullet')
    
    doc.add_heading('3.2 Data Quality and Integrity', 2)
    doc.add_paragraph(
        'Development data consists of 50,000 accounts originated between January 2021 and December 2022. '
        'Performance window: 12 months. Bad definition: 90+ days past due or charge-off within 12 months. '
        'Data quality checks include: missing value analysis (<2% missing), outlier detection and treatment, '
        'consistency checks across data sources, and temporal stability verification.'
    )
    
    doc.add_heading('3.3 Model Performance', 2)
    doc.add_paragraph('Development Sample Performance (50,000 accounts):')
    doc.add_paragraph('• KS Statistic: 0.45 (Excellent discrimination)', style='List Bullet')
    doc.add_paragraph('• Gini Coefficient: 0.65 (Good model performance)', style='List Bullet')
    doc.add_paragraph('• AUC-ROC: 0.825 (Excellent)', style='List Bullet')
    doc.add_paragraph('• Accuracy: 88.2%', style='List Bullet')
    doc.add_paragraph('• Precision: 45.3%', style='List Bullet')
    doc.add_paragraph('• Recall: 62.1%', style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('Out-of-Sample Validation (20,000 accounts):')
    doc.add_paragraph('• KS Statistic: 0.43 (Stable performance)', style='List Bullet')
    doc.add_paragraph('• Gini Coefficient: 0.63 (Minimal degradation)', style='List Bullet')
    doc.add_paragraph('• AUC-ROC: 0.815', style='List Bullet')
    
    doc.add_heading('3.4 Model Assumptions', 2)
    doc.add_paragraph('Key assumptions underlying the model include:')
    doc.add_paragraph('1. Historical relationships between predictors and default remain stable', style='List Number')
    doc.add_paragraph('2. Credit bureau data quality and reporting standards are maintained', style='List Number')
    doc.add_paragraph('3. Economic conditions remain within historical ranges observed in development data', style='List Number')
    doc.add_paragraph('4. Population characteristics and credit policies remain consistent', style='List Number')
    doc.add_paragraph('5. Definition of default (90+ DPD) remains unchanged', style='List Number')
    
    doc.add_paragraph()
    doc.add_paragraph(
        'These assumptions are monitored monthly through PSI/CSI analysis and quarterly through '
        'comprehensive model performance reviews.'
    )
    
    doc.add_heading('3.5 Ongoing Monitoring', 2)
    doc.add_paragraph('The model is subject to comprehensive ongoing monitoring:')
    doc.add_paragraph()
    doc.add_paragraph('Monthly Monitoring:')
    doc.add_paragraph('• Population Stability Index (PSI) - Trigger: PSI > 0.25', style='List Bullet')
    doc.add_paragraph('• Characteristic Stability Index (CSI) - Trigger: CSI > 0.25 for any variable', style='List Bullet')
    doc.add_paragraph('• Score distribution analysis', style='List Bullet')
    doc.add_paragraph('• Override rate tracking', style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('Quarterly Monitoring:')
    doc.add_paragraph('• KS Statistic and Gini Coefficient - Trigger: >15% degradation', style='List Bullet')
    doc.add_paragraph('• Discrimination power by score bands', style='List Bullet')
    doc.add_paragraph('• Calibration analysis', style='List Bullet')
    doc.add_paragraph('• Vintage analysis', style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('Annual Validation:')
    doc.add_paragraph('• Comprehensive independent validation', style='List Bullet')
    doc.add_paragraph('• Assumption testing', style='List Bullet')
    doc.add_paragraph('• Benchmark comparison', style='List Bullet')
    doc.add_paragraph('• Regulatory compliance review', style='List Bullet')
    
    # Model Variables
    doc.add_heading('4. Model Variables and Coefficients', 1)
    
    table = doc.add_table(rows=10, cols=4)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Variable Name'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Coefficient'
    hdr_cells[3].text = 'Significance'
    
    # Data rows
    variables = [
        ('Intercept', 'Constant', '2.1543', '***'),
        ('Credit Score', 'Continuous', '-0.0052', '***'),
        ('DTI Ratio', 'Continuous', '2.4567', '***'),
        ('Delinquencies_24m', 'Count', '0.3421', '***'),
        ('Employment_Length', 'Continuous', '-0.0087', '***'),
        ('Annual_Income', 'Continuous', '-0.0003', '***'),
        ('Credit_Utilization', 'Continuous', '0.5234', '***'),
        ('Num_Credit_Lines', 'Count', '-0.0234', '**'),
        ('Inquiries_6m', 'Count', '0.1456', '***'),
    ]
    
    for i, (var, vtype, coef, sig) in enumerate(variables, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = var
        row_cells[1].text = vtype
        row_cells[2].text = coef
        row_cells[3].text = sig
    
    # Model Limitations
    doc.add_heading('5. Model Limitations and Mitigants', 1)
    doc.add_paragraph('Known Limitations:')
    doc.add_paragraph('1. Limited performance data during economic downturns', style='List Number')
    doc.add_paragraph('2. Potential for credit bureau data quality issues', style='List Number')
    doc.add_paragraph('3. Does not capture all aspects of creditworthiness', style='List Number')
    
    doc.add_paragraph()
    doc.add_paragraph('Mitigating Controls:')
    doc.add_paragraph('• Stress testing under adverse scenarios', style='List Bullet')
    doc.add_paragraph('• Manual underwriter override capability', style='List Bullet')
    doc.add_paragraph('• Comprehensive monitoring program', style='List Bullet')
    doc.add_paragraph('• Regular model validation and recalibration', style='List Bullet')
    
    # Conclusion
    doc.add_heading('6. Conclusion', 1)
    doc.add_paragraph(
        'The US Unsecured Personal Loan Application Scorecard demonstrates strong predictive power '
        'and meets all SR 11-7 requirements for model risk management. The model is suitable for '
        'production use subject to the ongoing monitoring and validation program described in this document.'
    )
    
    # Save document
    doc.save(filename)
    print(f"✅ Created: {filename}")

def main():
    """Generate all sample files"""
    print("=" * 60)
    print("Generating Sample Data for Successful Model Validation")
    print("=" * 60)
    print()
    
    # Generate training data (2000 samples, 8% default rate)
    print("📊 Generating training data...")
    train_df = generate_features(2000, seed=42)
    train_df = generate_target_and_score(train_df, default_rate=0.08, seed=42)
    train_df.to_csv('successful_train.csv', index=False)
    print(f"✅ Created: successful_train.csv ({len(train_df)} rows)")
    print(f"   Default Rate: {train_df['target'].mean():.2%}")
    print(f"   Score Range: {train_df['score'].min()}-{train_df['score'].max()}")
    print()
    
    # Generate test data (1000 samples, similar distribution)
    print("📊 Generating test data...")
    test_df = generate_features(1000, seed=123)
    test_df = generate_target_and_score(test_df, default_rate=0.08, seed=123)
    test_df.to_csv('successful_test.csv', index=False)
    print(f"✅ Created: successful_test.csv ({len(test_df)} rows)")
    print(f"   Default Rate: {test_df['target'].mean():.2%}")
    print(f"   Score Range: {test_df['score'].min()}-{test_df['score'].max()}")
    print()
    
    # Generate OOT data (600 samples, slight shift)
    print("📊 Generating out-of-time data...")
    oot_df = generate_features(600, seed=456)
    oot_df = generate_target_and_score(oot_df, default_rate=0.09, seed=456)  # Slightly higher default
    oot_df.to_csv('successful_oot.csv', index=False)
    print(f"✅ Created: successful_oot.csv ({len(oot_df)} rows)")
    print(f"   Default Rate: {oot_df['target'].mean():.2%}")
    print(f"   Score Range: {oot_df['score'].min()}-{oot_df['score'].max()}")
    print()
    
    # Generate model documentation
    print("📄 Generating model documentation...")
    create_model_documentation('successful_model_documentation.docx')
    print()
    
    print("=" * 60)
    print("✅ All sample files generated successfully!")
    print("=" * 60)
    print()
    print("📁 Files created:")
    print("   1. successful_train.csv")
    print("   2. successful_test.csv")
    print("   3. successful_oot.csv")
    print("   4. successful_model_documentation.docx")
    print()
    print("🎯 Expected Validation Results:")
    print("   • KS Statistic: 0.40-0.50 (Excellent)")
    print("   • Gini Coefficient: 0.60-0.70 (Good)")
    print("   • PSI: < 0.10 (Stable)")
    print("   • CSI: < 0.10 (Stable)")
    print("   • Overall Status: PASS ✅")
    print()
    print("🚀 Next Steps:")
    print("   1. Open http://localhost:3002/")
    print("   2. Upload successful_model_documentation.docx (optional)")
    print("   3. Configure model (Application Scorecard, Unsecured Personal Loans)")
    print("   4. Upload the 3 CSV files")
    print("   5. Start validation and view results!")
    print()

if __name__ == '__main__':
    main()

# Made with Bob
